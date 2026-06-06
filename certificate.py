# نصب ماژول مورد نیاز (اگر قبلاً نصب نشده)
!pip install python-docx docx2pdf pandas openpyxl

import os
import pandas as pd
from docx import Document
from docx2pdf import convert

# ==========================
# تنظیمات اولیه
# ==========================

EXCEL_FILE = "winners.xlsx"  # نام فایل اکسل
TEMPLATE_FILE = "certificate.docx"  # نام فایل قالب Word
OUTPUT_FOLDER = "output"  # پوشه خروجی

# ==========================
# ایجاد پوشه خروجی
# ==========================

os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# ==========================
# تابع جایگزینی متن (نسخه ساده و مطمئن)
# ==========================

def replace_text(doc, placeholder, value):
    """
    جایگزینی متن در پاراگراف‌ها و جدول‌های سند Word
    """
    # جایگزینی در پاراگراف‌ها
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # جایگزینی مستقیم در متن پاراگراف
            paragraph.text = paragraph.text.replace(placeholder, str(value))
    
    # جایگزینی در جدول‌ها (اگر جدول وجود داشته باشد)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, str(value))

# ==========================
# خواندن فایل اکسل
# ==========================

try:
    df = pd.read_excel(EXCEL_FILE)
    print(f"✅ فایل اکسل با موفقیت خوانده شد.")
    print(f"📊 تعداد دانش‌آموزان: {len(df)}")
    print(f"📋 نام ستون‌ها: {list(df.columns)}")
    print("-" * 50)
except Exception as e:
    print(f"❌ خطا در خواندن فایل اکسل: {e}")
    print("لطفاً مطمئن شوید فایل winners.xlsx در مسیر جاری وجود دارد.")
    exit()

# ==========================
# تولید تشویق‌نامه‌ها
# ==========================

success_count = 0
error_count = 0

for index, row in df.iterrows():
    try:
        # خواندن نام و مدرسه از فایل اکسل
        name = str(row["NAME"]).strip()
        school = str(row["SCHOOL"]).strip()
        
        print(f"🔄 در حال پردازش {index + 1}: {name} - {school}")
        
        # باز کردن قالب Word
        doc = Document(TEMPLATE_FILE)
        
        # جایگزینی متغیرها
        replace_text(doc, "{{NAME}}", name)
        replace_text(doc, "{{SCHOOL}}", school)
        
        # ایجاد نام فایل ایمن (بدون کاراکترهای غیرمجاز)
        safe_name = name
        for char in ['/', '\\', ':', '*', '?', '"', '<', '>', '|']:
            safe_name = safe_name.replace(char, '-')
        
        # مسیرهای فایل‌ها
        temp_docx = os.path.join(OUTPUT_FOLDER, f"{safe_name}.docx")
        pdf_file = os.path.join(OUTPUT_FOLDER, f"{safe_name}.pdf")
        
        # ذخیره فایل Word موقت
        doc.save(temp_docx)
        
        # تبدیل به PDF
        convert(temp_docx, pdf_file)
        
        # حذف فایل Word موقت
        os.remove(temp_docx)
        
        print(f"   ✅ ساخته شد: {safe_name}.pdf")
        success_count += 1
        
    except Exception as e:
        print(f"   ❌ خطا برای {name}: {e}")
        error_count += 1
        continue

# ==========================
# گزارش نهایی
# ==========================

print("-" * 50)
print(f"📊 گزارش نهایی:")
print(f"   ✅ موفق: {success_count} تشویق‌نامه")
print(f"   ❌ ناموفق: {error_count} تشویق‌نامه")
print(f"   📁 مسیر خروجی: {OUTPUT_FOLDER}")
print("\n🎉 تمام تشویق‌نامه‌ها با موفقیت تولید شدند.")